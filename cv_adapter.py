import os
import re
import uuid
import docx2txt
from docx import Document
from collections import Counter

class CVAdapter:
    """
    Classe pour adapter un CV en fonction d'une description de poste.
    """
    
    def __init__(self, upload_folder, download_folder):
        """
        Initialise l'adaptateur de CV avec les dossiers de téléchargement et de stockage.
        
        Args:
            upload_folder (str): Chemin vers le dossier de téléchargement des CV
            download_folder (str): Chemin vers le dossier de stockage des CV adaptés
        """
        self.upload_folder = upload_folder
        self.download_folder = download_folder
        
        # Assurer que les dossiers existent
        os.makedirs(upload_folder, exist_ok=True)
        os.makedirs(download_folder, exist_ok=True)
    
    def extract_text_from_cv(self, cv_path):
        """
        Extrait le texte d'un fichier CV au format docx.
        
        Args:
            cv_path (str): Chemin vers le fichier CV
            
        Returns:
            str: Texte extrait du CV
        """
        try:
            return docx2txt.process(cv_path)
        except Exception as e:
            print(f"Erreur lors de l'extraction du texte du CV: {e}")
            return ""
    
    def extract_keywords_from_job_description(self, job_description):
        """
        Extrait les mots-clés importants d'une description de poste.
        
        Args:
            job_description (str): Description du poste
            
        Returns:
            list: Liste des mots-clés importants
        """
        # Nettoyage du texte
        job_description = job_description.lower()
        
        # Suppression des mots vides (stopwords) en français
        stopwords = [
            "le", "la", "les", "un", "une", "des", "et", "ou", "de", "du", "au", "aux",
            "ce", "cette", "ces", "mon", "ton", "son", "notre", "votre", "leur",
            "pour", "par", "sur", "sous", "dans", "avec", "sans", "en", "à", "qui",
            "que", "quoi", "dont", "où", "comment", "pourquoi", "quand", "est", "sont",
            "sera", "seront", "été", "avoir", "être", "faire", "plus", "moins", "très"
        ]
        
        # Extraction des mots
        words = re.findall(r'\b\w{3,}\b', job_description)
        
        # Filtrage des mots vides et comptage des occurrences
        filtered_words = [word for word in words if word not in stopwords]
        word_counts = Counter(filtered_words)
        
        # Sélection des mots-clés les plus fréquents (au moins 2 occurrences)
        keywords = [word for word, count in word_counts.items() if count >= 2]
        
        # Ajout des mots de plus de 6 caractères (potentiellement importants)
        long_words = [word for word in filtered_words if len(word) > 6 and word not in keywords]
        
        # Combinaison des mots-clés
        all_keywords = list(set(keywords + long_words))
        
        return all_keywords
    
    def adapt_cv(self, cv_path, job_description):
        """
        Adapte un CV en fonction d'une description de poste.
        
        Args:
            cv_path (str): Chemin vers le fichier CV
            job_description (str): Description du poste
            
        Returns:
            str: Nom du fichier CV adapté
        """
        # Extraire les mots-clés de la description du poste
        keywords = self.extract_keywords_from_job_description(job_description)
        
        # Charger le document Word original pour le modifier
        doc = Document(cv_path)
        
        # Adapter le CV en mettant en évidence les compétences pertinentes
        for paragraph in doc.paragraphs:
            for keyword in keywords:
                if keyword in paragraph.text.lower():
                    # Mettre en gras les mots-clés trouvés
                    text = paragraph.text
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    
                    # Sauvegarde du texte original
                    original_text = paragraph.text
                    
                    # Effacer le contenu du paragraphe
                    paragraph.clear()
                    
                    # Trouver toutes les occurrences du mot-clé
                    matches = list(pattern.finditer(original_text))
                    
                    # Si aucune correspondance, continuer
                    if not matches:
                        paragraph.add_run(original_text)
                        continue
                    
                    # Ajouter le texte avec les mots-clés en gras
                    last_end = 0
                    for match in matches:
                        # Ajouter le texte avant le mot-clé
                        if match.start() > last_end:
                            paragraph.add_run(original_text[last_end:match.start()])
                        
                        # Ajouter le mot-clé en gras
                        paragraph.add_run(original_text[match.start():match.end()]).bold = True
                        
                        last_end = match.end()
                    
                    # Ajouter le reste du texte
                    if last_end < len(original_text):
                        paragraph.add_run(original_text[last_end:])
        
        # Générer un nom de fichier unique pour le CV adapté
        output_filename = f"cv_adapte_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(self.download_folder, output_filename)
        
        # Sauvegarder le document modifié
        doc.save(output_path)
        
        return output_filename
    
    def analyze_job_description(self, job_description):
        """
        Analyse une description de poste pour extraire des informations utiles.
        
        Args:
            job_description (str): Description du poste
            
        Returns:
            dict: Informations extraites de la description du poste
        """
        # Extraction des compétences techniques potentielles
        tech_skills_pattern = r'compétences\s+techniques.*?:(.*?)(?:\n\n|\Z)'
        tech_skills_match = re.search(tech_skills_pattern, job_description, re.IGNORECASE | re.DOTALL)
        tech_skills = []
        if tech_skills_match:
            tech_skills_text = tech_skills_match.group(1)
            tech_skills = [skill.strip() for skill in re.split(r'[,;•]', tech_skills_text) if skill.strip()]
        
        # Extraction du niveau d'expérience requis
        experience_pattern = r'(\d+)[+]?\s+ans?\s+d\'expérience'
        experience_match = re.search(experience_pattern, job_description, re.IGNORECASE)
        experience_years = experience_match.group(1) if experience_match else None
        
        # Extraction du niveau d'études requis
        education_patterns = [
            r'(bac\s*[+]\s*\d+)',
            r'(master|licence|doctorat)',
            r'(diplôme\s+d\'ingénieur)'
        ]
        education = None
        for pattern in education_patterns:
            match = re.search(pattern, job_description, re.IGNORECASE)
            if match:
                education = match.group(1)
                break
        
        # Résultats de l'analyse
        analysis = {
            'keywords': self.extract_keywords_from_job_description(job_description),
            'technical_skills': tech_skills,
            'experience_years': experience_years,
            'education': education
        }
        
        return analysis
