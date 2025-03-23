import os
import docx
import uuid
import tempfile
from docx import Document

# Créer un CV de test simple
def create_test_cv():
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Jean Dupont', 0)
    
    # Ajouter des informations de contact
    doc.add_paragraph('Email: jean.dupont@email.com | Téléphone: 06 12 34 56 78 | Paris, France')
    
    # Ajouter une section compétences
    doc.add_heading('Compétences', level=1)
    skills = doc.add_paragraph()
    skills.add_run('Développement web: ').bold = True
    skills.add_run('HTML, CSS, JavaScript, React, Node.js\n')
    skills.add_run('Programmation: ').bold = True
    skills.add_run('Python, Java, C++\n')
    skills.add_run('Base de données: ').bold = True
    skills.add_run('MySQL, MongoDB, PostgreSQL\n')
    skills.add_run('Outils: ').bold = True
    skills.add_run('Git, Docker, AWS, Azure')
    
    # Ajouter une section expérience
    doc.add_heading('Expérience professionnelle', level=1)
    
    # Expérience 1
    exp1 = doc.add_paragraph()
    exp1.add_run('Développeur Full Stack, TechCorp').bold = True
    exp1.add_run(' | Janvier 2022 - Présent\n')
    exp1.add_run('• Développement d\'applications web avec React et Node.js\n')
    exp1.add_run('• Mise en place d\'une architecture microservices\n')
    exp1.add_run('• Optimisation des performances et de la sécurité des applications')
    
    # Expérience 2
    exp2 = doc.add_paragraph()
    exp2.add_run('Développeur Frontend, WebSolutions').bold = True
    exp2.add_run(' | Mars 2020 - Décembre 2021\n')
    exp2.add_run('• Création d\'interfaces utilisateur avec HTML, CSS et JavaScript\n')
    exp2.add_run('• Intégration d\'API RESTful\n')
    exp2.add_run('• Collaboration avec les designers UX/UI')
    
    # Ajouter une section formation
    doc.add_heading('Formation', level=1)
    edu = doc.add_paragraph()
    edu.add_run('Master en Informatique, Université de Paris').bold = True
    edu.add_run(' | 2018 - 2020\n')
    edu.add_run('Licence en Informatique, Université de Lyon | 2015 - 2018')
    
    # Ajouter une section langues
    doc.add_heading('Langues', level=1)
    lang = doc.add_paragraph()
    lang.add_run('Français: ').bold = True
    lang.add_run('Langue maternelle\n')
    lang.add_run('Anglais: ').bold = True
    lang.add_run('Courant (TOEIC 850)\n')
    lang.add_run('Espagnol: ').bold = True
    lang.add_run('Intermédiaire')
    
    # Sauvegarder le document
    test_cv_path = os.path.join('uploads', f'test_cv_{uuid.uuid4().hex}.docx')
    doc.save(test_cv_path)
    
    return test_cv_path

# Créer des descriptions de poste de test
def create_test_job_descriptions():
    descriptions = []
    
    # Description 1: Développeur web
    dev_web = """
    Développeur Web Full Stack (H/F)
    
    Notre entreprise recherche un développeur web full stack expérimenté pour rejoindre notre équipe de développement.
    
    Responsabilités:
    • Développer des applications web performantes et réactives
    • Travailler avec les technologies frontend et backend
    • Collaborer avec les designers UX/UI
    • Participer aux revues de code et aux tests
    
    Compétences techniques:
    • Maîtrise de HTML, CSS et JavaScript
    • Expérience avec React.js et Node.js
    • Connaissance des bases de données SQL et NoSQL
    • Expérience avec Git et les méthodologies agiles
    
    Profil recherché:
    • 3+ ans d'expérience en développement web
    • Diplôme en informatique ou équivalent
    • Capacité à travailler en équipe
    • Anglais professionnel
    
    Nous offrons:
    • Un environnement de travail stimulant
    • Des projets variés et innovants
    • Une rémunération compétitive
    • Des possibilités d'évolution
    """
    
    # Description 2: Data Scientist
    data_scientist = """
    Data Scientist (H/F)
    
    Nous recherchons un data scientist talentueux pour analyser nos données et développer des modèles prédictifs.
    
    Responsabilités:
    • Analyser de grands ensembles de données
    • Développer des modèles de machine learning
    • Présenter les résultats aux parties prenantes
    • Collaborer avec les équipes produit et ingénierie
    
    Compétences techniques:
    • Maîtrise de Python et des bibliothèques de data science (pandas, scikit-learn)
    • Expérience avec les techniques de machine learning
    • Connaissance des bases de données SQL
    • Compétences en visualisation de données
    
    Profil recherché:
    • Master ou PhD en informatique, mathématiques ou domaine connexe
    • 2+ ans d'expérience en data science
    • Solides compétences analytiques
    • Capacité à communiquer des concepts complexes
    
    Nous offrons:
    • Un environnement de travail collaboratif
    • Des défis techniques stimulants
    • Une rémunération attractive
    • Des avantages sociaux compétitifs
    """
    
    descriptions.append(dev_web)
    descriptions.append(data_scientist)
    
    return descriptions

# Fonction principale de test
def run_tests():
    print("=== Début des tests de l'application d'adaptation de CV ===")
    
    # Créer un CV de test
    test_cv_path = create_test_cv()
    print(f"CV de test créé: {test_cv_path}")
    
    # Créer des descriptions de poste de test
    job_descriptions = create_test_job_descriptions()
    print(f"Descriptions de poste de test créées: {len(job_descriptions)}")
    
    # Importer la classe CVAdapter
    from cv_adapter import CVAdapter
    
    # Initialiser l'adaptateur de CV
    adapter = CVAdapter('uploads', 'downloads')
    
    # Tester l'adaptation du CV pour chaque description de poste
    for i, description in enumerate(job_descriptions):
        print(f"\n=== Test {i+1}: Adaptation pour '{description.splitlines()[1].strip()}' ===")
        
        # Adapter le CV
        try:
            adapted_cv_filename = adapter.adapt_cv(test_cv_path, description)
            print(f"CV adapté créé: {adapted_cv_filename}")
            
            # Vérifier que le fichier existe
            adapted_cv_path = os.path.join('downloads', adapted_cv_filename)
            if os.path.exists(adapted_cv_path):
                print("✓ Le fichier CV adapté existe")
                
                # Vérifier que le contenu a été modifié (mots en gras)
                doc = Document(adapted_cv_path)
                bold_words_count = 0
                
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.bold:
                            bold_words_count += 1
                
                print(f"✓ Nombre de segments en gras: {bold_words_count}")
                
                # Analyser la description de poste
                analysis = adapter.analyze_job_description(description)
                print(f"✓ Mots-clés extraits: {len(analysis['keywords'])}")
                print(f"✓ Compétences techniques identifiées: {len(analysis['technical_skills'])}")
                
            else:
                print("✗ Le fichier CV adapté n'existe pas")
        
        except Exception as e:
            print(f"✗ Erreur lors de l'adaptation du CV: {str(e)}")
    
    print("\n=== Tests terminés ===")

if __name__ == "__main__":
    run_tests()
